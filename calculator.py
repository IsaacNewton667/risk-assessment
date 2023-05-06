import math
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from pathlib import Path
import os


e = math.e

variant = 0
risk = 0
count = 0
position = 0

alfa = []
beta = []
Tmesh = []
Tdiag = []
Tvost = []
Tzad = []

alfaTime = []
betaTime = []
TmeshTime = []
TdiagTime = []
TvostTime = []
TzadTime = []

TzadTimeReliability = "year"
TzadTimeRequirements = "year"

axisX = []
axisY = []

riskReliability = 0
riskRequirements = 0


def toFixed(numObj, digits=4):
    return f"{numObj:.{digits}f}"


def add_parameters_for_reliability(alfaAdd, betaAdd, TmeshAdd, TdiagAdd, TvostAdd, TzadAdd, alfaTimeAdd, betaTimeAdd,
                                   TmeshTimeAdd, TdiagTimeAdd, TvostTimeAdd, TzadTimeAdd):
    global alfa
    global beta
    global Tmesh
    global Tdiag
    global Tvost
    global Tzad

    global alfaTime
    global betaTime
    global TmeshTime
    global TdiagTime
    global TvostTime
    global TzadTime

    alfa.append(alfaAdd)
    beta.append(betaAdd)
    Tmesh.append(TmeshAdd)
    Tdiag.append(TdiagAdd)
    Tvost.append(TvostAdd)
    Tzad.append(TzadAdd)

    alfaTime.append(alfaTimeAdd)
    betaTime.append(betaTimeAdd)
    TmeshTime.append(TmeshTimeAdd)
    TdiagTime.append(TdiagTimeAdd)
    TvostTime.append(TvostTimeAdd)
    TzadTime.append(TzadTimeAdd)

    global count
    count = count + 1
    global axisX
    axisX.append(count)


def calc_time(data, time, Tzad):
    # if time != Tzad
    if Tzad == "year" and time != Tzad:
        match time:
            case "month":
                data = data / 12
                return data
            case "week":
                data = data / 52.1429
                return data
            case "day":
                data = data / 365
                return data
            case "hour":
                data = data / 8760
                return data
            case "min":
                data = data / 525600
                return data
            case "sec":
                data = data / (3.154 * 10 ** 7)
                return data
            case "year":
                return data
    elif Tzad == "month" and time != Tzad:
        match time:
            case "month":
                return data
            case "week":
                data = data * 4.34
                return data
            case "day":
                data = data * 30
                return data
            case "hour":
                data = data * 730
                return data
            case "min":
                data = data * 43800
                return data
            case "sec":
                data = data * (2.628 * 10 ** 6)
                return data
            case "year":
                data = data / 12
                return data
    else:
        return data


# выбор варианта
def variants(Tzad, Tmesh, Tdiag):
    if Tzad < (Tmesh + Tdiag):
        return 1
    else:
        return 2


# вычисление риска
def calc_reliability_risk_variant_1(alfa, beta, Tzad):
      #  risk = 1 - ((alfa - beta ** (-1)) ** (-1) * (alfa * e ** (-Tzad / beta) - beta ** (-1) * e ** (-alfa * Tzad)))

      if alfa != beta ** (-1):
          RiskPartOne = (alfa - beta ** (-1)) ** (-1)
          RiskPartTwo = alfa * e ** (-Tzad / beta)
          RiskPartThree = beta ** (-1) * e ** (-alfa * Tzad)

          risk = 1 - (RiskPartOne * (RiskPartTwo - RiskPartThree))

          return risk

      elif alfa == beta ** (-1):
          RiskPartOne = e ** (-alfa * Tzad)
          RiskPartTwo = 1 + alfa * Tzad

          risk = 1 - (RiskPartOne * RiskPartTwo)
          return risk
      else:
          return 0


def calc_reliability_risk_variant_2(alfa, beta, Tzad, Tmesh, Tdiag):
    summTmeshTdiag = Tmesh + Tdiag

    N = math.floor(Tzad / (summTmeshTdiag))
    Tost = Tzad - N * (summTmeshTdiag)

    if alfa != beta ** (-1):
        PseredPartOne = (alfa - beta ** (-1)) ** (-1)
        PseredPartTwo = alfa * e ** (-summTmeshTdiag / beta)
        PseredPartThree = beta ** (-1) * e ** (-alfa * summTmeshTdiag)

        Psered = (PseredPartOne * (PseredPartTwo - PseredPartThree)) ** (N)

        #   Psered = ((alfa - beta**(-1))**(-1) * (alfa * e**(-summTmeshTdiag/beta) - beta**(-1) * e**(-alfa*summTmeshTdiag)))**(N)

        PkonPartOne = (alfa - beta ** (-1)) ** (-1)
        PkonPartTwo = alfa * e ** (-Tost / beta)
        PkonPartThree = beta ** (-1) * e ** (-alfa * Tost)

        Pkon = PkonPartOne * (PkonPartTwo - PkonPartThree)

        #   Pkon = (alfa - beta**(-1))**(1) * (alfa * e**(-Tost/beta) - beta**(-1) * e**(-alfa*Tost))

        risk = 1 - (Psered * Pkon)
        return risk

    elif alfa == beta ** (-1):
        PseredPartOne = e ** (-alfa * summTmeshTdiag)
        PseredPartTwo = 1 + alfa * summTmeshTdiag

        Psered = (PseredPartOne * PseredPartTwo) ** (N)

        PkonPartOne = e ** (-alfa * Tost)
        PkonPartTwo = 1 + alfa * Tost

        Pkon = PkonPartOne * PkonPartTwo

        risk = 1 - (Psered * Pkon)
        return risk
    else:
        return 0


def calc_reliability_risk(alfa, beta, Tzad, Tmesh, Tdiag):
    if variant == 1:
        risk = calc_reliability_risk_variant_1(alfa, beta, Tzad)
    elif variant == 2:
        risk = calc_reliability_risk_variant_2(alfa, beta, Tzad, Tmesh, Tdiag)


def start_calc_reliability_risk():
    global position
    global count
    global axisY
    global axisX

    global alfa
    global beta
    global Tmesh
    global Tdiag
    global Tvost
    global Tzad

    global alfaTime
    global betaTime
    global TmeshTime
    global TdiagTime
    global TvostTime
    global TzadTime

    global TzadTimeReliability

    while position < count:
        alfa[position] = float(alfa[position])
        beta[position] = int(beta[position])
        Tmesh[position] = int(Tmesh[position])
        Tdiag[position] = int(Tdiag[position])
        Tvost[position] = int(Tvost[position])
        Tzad[position] = int(Tzad[position])

        alfa[position] = calc_time(alfa[position], alfaTime[position], TzadTime[position])
        beta[position] = calc_time(beta[position], betaTime[position], TzadTime[position])
        Tmesh[position] = calc_time(Tmesh[position], TmeshTime[position], TzadTime[position])
        Tdiag[position] = calc_time(Tdiag[position], TdiagTime[position], TzadTime[position])
        Tvost[position] = calc_time(Tvost[position], TvostTime[position], TzadTime[position])

        TzadTimeReliability = TzadTime[position]

        variant = variants(Tzad[position], Tmesh[position], Tdiag[position])

        if variant == 1:
            risk = calc_reliability_risk_variant_1(alfa[position], beta[position], Tzad[position])
        elif variant == 2:
            risk = calc_reliability_risk_variant_2(alfa[position], beta[position], Tzad[position], Tmesh[position],
                                                   Tdiag[position])

        axisY.append(risk)
        position = position + 1

    return axisY


def start_calc_requirements_risk():
    global position
    global count
    global axisY
    global axisX

    global alfa
    global beta
    global Tmesh
    global Tdiag
    global Tvost
    global Tzad

    global alfaTime
    global betaTime
    global TmeshTime
    global TdiagTime
    global TvostTime
    global TzadTime

    global TzadTimeRequirements

    while position < count:
        alfa[position] = float(alfa[position])
        beta[position] = int(beta[position])
        Tmesh[position] = int(Tmesh[position])
        Tdiag[position] = int(Tdiag[position])
        Tvost[position] = int(Tvost[position])
        Tzad[position] = int(Tzad[position])

        alfa[position] = calc_time(alfa[position], alfaTime[position], TzadTime[position])
        beta[position] = calc_time(beta[position], betaTime[position], TzadTime[position])
        Tmesh[position] = calc_time(Tmesh[position], TmeshTime[position], TzadTime[position])
        Tdiag[position] = calc_time(Tdiag[position], TdiagTime[position], TzadTime[position])
        Tvost[position] = calc_time(Tvost[position], TvostTime[position], TzadTime[position])

        TzadTimeRequirements = TzadTime[position]

        variant = variants(Tzad[position], Tmesh[position], Tdiag[position])

        if variant == 1:
            risk = calc_reliability_risk_variant_1(alfa[position], beta[position], Tzad[position])
        elif variant == 2:
            risk = calc_reliability_risk_variant_2(alfa[position], beta[position], Tzad[position], Tmesh[position],
                                                   Tdiag[position])

        axisY.append(risk)
        position = position + 1

    return axisY


def clear_parameters():
    global position
    global count
    global variant
    global risk
    global axisY
    global axisX

    global alfa
    global beta
    global Tmesh
    global Tdiag
    global Tvost
    global Tzad

    global alfaTime
    global betaTime
    global TmeshTime
    global TdiagTime
    global TvostTime
    global TzadTime

    alfa.clear()
    beta.clear()
    Tmesh.clear()
    Tdiag.clear()
    Tvost.clear()
    Tzad.clear()

    alfaTime.clear()
    betaTime.clear()
    TmeshTime.clear()
    TdiagTime.clear()
    TvostTime.clear()
    TzadTime.clear()

    axisX.clear()
    axisY.clear()

    variant = 0
    risk = 0
    count = 0
    position = 0

    return None


def create_report_reliability():
    global axisX
    global axisY
    global count
    global riskReliability

    summY = 0

    pathImage = "images/figure.png"
    pathReport = "static/reports/report.docx"

    oldImage = str(Path("images", "figure.png").absolute())
    oldReport = str(Path("static", "reports", "report.docx").absolute())

    if os.path.isfile(oldImage):
        os.remove(oldImage)

    if os.path.isfile(oldReport):
        os.remove(oldReport)

    # создание графика
    fig = plt.figure()

    #    x = [1, 2, 3, 4, 5, 6, 7, 8, 9]

    # в список Y записываешь сами результаты оценки рисков, на графике самостоятельно выведется шкала на оси Y

    #    y = [0.067, 0.020, 0.020, 0.020, 0.015, 0.026, 0.062, 0.027, 0.186]
    count = count + 1
    axisX.append(count)
    reliabilityAxisX = axisX.copy()
    reliabilityAxisY = axisY.copy()

    reliabilityAxisX.remove(count)
    reliabilityAxisX.append("За все \n действия")

    for item in reliabilityAxisY:
        summY = summY + item

    riskReliability = summY

    reliabilityAxisYElements = reliabilityAxisY.copy()
    reliabilityAxisY.append(summY)
    # print(axisX)
    plt.bar(axisX, reliabilityAxisY)

    plt.xticks(axisX, reliabilityAxisX)
    # plt.xticks(axisX, ['1', '2', '3', '4', '5', '6', '7', '8', 'За все \n действия'])

    plt.show()
    fig.savefig(oldImage, dpi=100)

    # создание репорта

    # создание документа
    report = Document()

    # задаем стиль текста по умолчанию
    style = report.styles['Normal']
    # название шрифта
    style.font.name = 'Times New Roman'
    # размер шрифта
    style.font.size = Pt(14)

    # добавление заголовка
    head = report.add_paragraph('Результаты вычислений количественных показателей рисков нарушения надежности, требований и интегрального риска\n')
    # выравнивание заголовка посередине
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # добавление параграф
    report.add_paragraph(f'Расчет количественных показателей оценки рисков процесса управления решениями показал, что за весь указанный период риск нарушения надежности составляет {toFixed(riskReliability)}. Результаты расчетов риска для каждого из элементов:\n')

    idElement = 1
    for item in reliabilityAxisYElements:
        report.add_paragraph(f'{idElement} элемент - {toFixed(item)}')
        idElement = idElement + 1

    # добавить картинку
    print(oldImage)
    picture = report.add_picture(oldImage)

    # сохранить репорт
    report.save(oldReport)

    return None

"""
def create_report():
    global axisX
    global axisY
    global count

    existReport = "static/reports/report.docx"

    summY = 0

    # создание графика
    fig = plt.figure()

    #    x = [1, 2, 3, 4, 5, 6, 7, 8, 9]

    # в список Y записываешь сами результаты оценки рисков, на графике самостоятельно выведется шкала на оси Y

    #    y = [0.067, 0.020, 0.020, 0.020, 0.015, 0.026, 0.062, 0.027, 0.186]
    count = count + 1
    axisX.append(count)
    reliabilityAxisX = axisX.copy()
    reliabilityAxisY = axisY.copy()

    reliabilityAxisX.remove(count)
    reliabilityAxisX.append("За все \n действия")

    for item in reliabilityAxisY:
        summY = summY + item

    reliabilityAxisY.append(summY)
    print(axisX)
    plt.bar(axisX, reliabilityAxisY)

    plt.xticks(axisX, reliabilityAxisX)
    # plt.xticks(axisX, ['1', '2', '3', '4', '5', '6', '7', '8', 'За все \n действия'])

    plt.show()
    fig.savefig('images/figure.png', dpi=100)

    # создание репорта

    # создание документа
    report = Document(existReport)

    # задаем стиль текста по умолчанию
    style = report.styles['Normal']
    # название шрифта
    style.font.name = 'Times New Roman'
    # размер шрифта
    style.font.size = Pt(14)

    # добавление заголовка
    head = report.add_paragraph('ОЦЕНКА РИСКОВ \n')
    # выравнивание заголовка посередине
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # добавление параграф
    report.add_paragraph('Оценка рисков \n')

    # добавить картинку
    picture = report.add_picture('images/figure.png')

    # сохранить репорт
    report.save(oldReport)

    return None
"""


def add_report_requirements():
    global axisX
    global axisY
    global count
    global riskRequirements
    summY = 0

    oldImage = str(Path("images", "figure.png").absolute())
    oldReport = str(Path("static", "reports", "report.docx").absolute())

    if os.path.isfile(oldImage):
        os.remove(oldImage)

    # создание графика
    fig = plt.figure()

    #    x = [1, 2, 3, 4, 5, 6, 7, 8, 9]

    # в список Y записываешь сами результаты оценки рисков, на графике самостоятельно выведется шкала на оси Y

    #    y = [0.067, 0.020, 0.020, 0.020, 0.015, 0.026, 0.062, 0.027, 0.186]
    count = count + 1
    axisX.append(count)
    requirementsAxisX = axisX.copy()
    requirementsAxisY = axisY.copy()

    requirementsAxisX.remove(count)
    requirementsAxisX.append("За все \n действия")

    for item in requirementsAxisY:
        summY = summY + item

    riskRequirements = summY

    requirementsAxisYElements = requirementsAxisY.copy()
    requirementsAxisY.append(summY)
    print(axisX)
    plt.bar(axisX, requirementsAxisY)

    plt.xticks(axisX, requirementsAxisX)
    # plt.xticks(axisX, ['1', '2', '3', '4', '5', '6', '7', '8', 'За все \n действия'])

    plt.show()
    fig.savefig(oldImage, dpi=100)

    # создание репорта

    # создание документа
    report = Document(oldReport)

    # задаем стиль текста по умолчанию
    style = report.styles['Normal']
    # название шрифта
    style.font.name = 'Times New Roman'
    # размер шрифта
    style.font.size = Pt(14)

    # добавление параграф
    report.add_paragraph('\n')
    report.add_paragraph(
        f'Расчет количественных показателей оценки рисков процесса управления решениями показал, что за весь указанный период риск нарушения требований составляет {toFixed(riskRequirements)}. Результаты расчетов риска для каждого из элементов:\n')

    idElement = 1
    for item in requirementsAxisYElements:
        report.add_paragraph(f'{idElement} элемент - {toFixed(item)}')
        idElement = idElement + 1

    # добавить картинку
    picture = report.add_picture(oldImage)

    # сохранить репорт
    report.save(oldReport)

    return None

def add_integral_risk():
    global riskReliability
    global riskRequirements
    global TzadTimeRequirements
    global TzadTimeReliability

    if TzadTimeReliability == "year" and TzadTimeRequirements == "month":
        riskReliability = riskReliability / 12
    elif TzadTimeRequirements == "year" and TzadTimeReliability == "month":
        riskRequirements = riskRequirements / 12

    # oldImage = str(Path("images", "figure.png").absolute())
    oldReport = str(Path("static", "reports", "report.docx").absolute())

    integralRisk = 1 - (1 - riskReliability) * (1 - riskRequirements)

    # создание документа
    report = Document(oldReport)

    # задаем стиль текста по умолчанию
    style = report.styles['Normal']
    # название шрифта
    style.font.name = 'Times New Roman'
    # размер шрифта
    style.font.size = Pt(14)

    # добавление параграф
    report.add_paragraph('\n')
    report.add_paragraph(f'Интегральный риск: \n\n R(интегр) = 1 - (1 - R(надежн)) * (1 - R(наруш)) = {toFixed(integralRisk)}')

    if integralRisk <= 0.05:
        report.add_paragraph(f'Интегральный риск равен {integralRisk}, этот результат меньше допустимого уровня 0,05. Это подтверждает сбалансированность планируемых к применению или применяемых технических решений с точкки зрения достижения целей системной инженерии.')
    else:
        report.add_paragraph(
            f'Интегральный риск равен {integralRisk}, этот результат больше допустимого уровня 0,05. Это является подтверждением того, что планируемые к применению или применяемые технические решения с точки зрения достижения целей системной инженерии не являются сбалансированным.')
    # сохранить репорт
    report.save(oldReport)

    return None